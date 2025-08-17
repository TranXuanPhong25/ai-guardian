import os
import tempfile
import requests
from typing import Optional
import logging
from docx import Document as DocxDocument
import pandas as pd
import torch
import fitz  # PyMuPDF
from docling_core.types.doc.document import DocTagsDocument, DoclingDocument
from transformers import AutoProcessor, AutoModelForVision2Seq
from transformers.image_utils import load_image
from urllib.parse import urlparse

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileExtractorService:
    def __init__(self):
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        torch.set_default_device(self.device)
        self.processor = AutoProcessor.from_pretrained("ds4sd/SmolDocling-256M-preview")
        self.model = AutoModelForVision2Seq.from_pretrained("ds4sd/SmolDocling-256M-preview").to(self.device)

    def _smoldocling_extract(self, image_path: str) -> str:
        """Extract content from an image using SmolDocling."""
        try:
            image = load_image(image_path)
            messages = [
                {"role": "user", "content": [
                    {"type": "image"},
                    {"type": "text", "text": "Convert this page to docling."}
                ]}
            ]
            prompt = self.processor.apply_chat_template(messages, add_generation_prompt=True)
            inputs = self.processor(text=prompt, images=[image], return_tensors="pt").to(self.device)
            generated_ids = self.model.generate(**inputs, max_new_tokens=8192)
            prompt_len = inputs.input_ids.shape[1]
            trimmed = generated_ids[:, prompt_len:]
            doctags = self.processor.batch_decode(trimmed, skip_special_tokens=False)[0].lstrip()
            doctags_doc = DocTagsDocument.from_doctags_and_image_pairs([doctags], [image])
            doc = DoclingDocument.load_from_doctags(doctags_doc)
            return doc.export_to_markdown()
        except Exception as e:
            logger.error(f"Error in smoldocling_extract: {str(e)}")
            return ""

    def _download_file_from_url(self, file_url: str, file_extension: str) -> str:
        """Download file from URL to temporary local file."""
        try:
            response = requests.get(file_url)
            response.raise_for_status()
            
            # Create temporary file with proper extension
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                temp_file.write(response.content)
                return temp_file.name
        except Exception as e:
            logger.error(f"Error downloading file from URL {file_url}: {str(e)}")
            raise

    def extract_text(self, file_path: str, file_format: str = ".png") -> Optional[str]:
        """Extract text from a file based on its extension. 
        file_path can be either a local path or a URL."""
        logger.info("extracting "+ file_path)
        
        # Check if file_path is a URL
        is_url = urlparse(file_path).scheme in ('http', 'https')
        temp_file_path = None
        
        try:
            if is_url:
                # Download file to temporary location
                temp_file_path = self._download_file_from_url(file_path, file_format)
                actual_file_path = temp_file_path
            else:
                actual_file_path = file_path
                
            extension = file_format
            text = ""

            if extension in ['.png', '.jpg', '.jpeg']:
                text = self._smoldocling_extract(actual_file_path)
            elif extension == '.pdf':
                doc = fitz.open(actual_file_path)
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    page_text = page.get_text("text")
                    text += page_text + "\n\n"
                    for img_index, img in enumerate(page.get_images(full=True)):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        temp_img_path = f"/tmp/pdf_page_{page_num}_img_{img_index}.png"
                        with open(temp_img_path, "wb") as f:
                            f.write(image_bytes)
                        page_text = self._smoldocling_extract(temp_img_path)
                        text += page_text + "\n\n"
                        os.remove(temp_img_path)
                doc.close()

            elif extension == '.docx':
                doc = DocxDocument(actual_file_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = "\t".join(cell.text.strip() for cell in row.cells)
                        text += row_text + "\n"

            elif extension in ['.xls', '.xlsx', '.csv']:
                if extension == '.csv':
                    df = pd.read_csv(actual_file_path)
                    text += df.to_string(index=False) + "\n\n"
                else:
                    sheets = pd.read_excel(actual_file_path, sheet_name=None)
                    for sheet_name, sheet_data in sheets.items():
                        text += f"--- Sheet: {sheet_name} ---\n"
                        text += sheet_data.to_string(index=False) + "\n\n"

            else:
                raise ValueError(f"Unsupported file format: {extension}")

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
        finally:
            # Clean up temporary file if it was created
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.remove(temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to remove temporary file {temp_file_path}: {str(e)}")

file_extractor_service = FileExtractorService()